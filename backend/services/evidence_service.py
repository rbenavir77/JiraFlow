import os
import cv2
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import datetime
import tkinter as tk
from tkinter import filedialog

import shutil

class EvidenceService:
    def __init__(self, template_path=None):
        # Apuntar directamente al archivo con extensión simple
        self.template_path = os.path.abspath(os.path.join("backend", "templates", "template_evidencia.docx"))
        
        print(f"[EvidenceService] Plantilla configurada en: {self.template_path}")
        if not os.path.exists(self.template_path):
            print(f"[EvidenceService] ADVERTENCIA: No se encontró la plantilla 'template_evidencia.docx'.")

    def pick_directory(self):
        """Abre un diálogo nativo de Windows para seleccionar una carpeta."""
        root = tk.Tk()
        root.withdraw() # Ocultar la ventana principal de tkinter
        root.attributes('-topmost', True) # Asegurar que el diálogo aparezca al frente
        directory = filedialog.askdirectory()
        root.destroy()
        return directory

    def _get_temp_dir(self):
        """Calcula una ruta temporal dinámica usando la carpeta oficial del sistema operativo."""
        import tempfile
        # Creamos una subcarpeta específica para nuestro aplicativo en la zona temporal del sistema
        temp_base = os.path.join(tempfile.gettempdir(), "JiraFlow_Evidence_Temp")
        return os.path.abspath(temp_base)

    def extract_frames_from_video(self, video_path, num_frames=3):
        """Extrae N capturas de un video y las guarda en una carpeta temporal central."""
        video_path = os.path.abspath(video_path)
        output_folder = self._get_temp_dir()
            
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
            
        print(f"[EvidenceService] Procesando video: {video_path}")
        
        # Prefijo para rutas largas en Windows
        win_path = video_path
        if len(win_path) > 240 and os.name == 'nt':
            win_path = "\\\\?\\" + win_path

        cap = cv2.VideoCapture(win_path)
        if not cap.isOpened():
            print(f"[EvidenceService] ERROR: No se pudo abrir el video.")
            return []

        total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        if total_frames <= 0:
            cap.release()
            return []

        extracted_paths = []
        import unicodedata
        stem = Path(video_path).stem
        video_stem = unicodedata.normalize('NFKD', stem).encode('ascii', 'ignore').decode('ascii')
        video_stem = "".join([c for c in video_stem if c.isalnum() or c in (' ', '_', '-')]).strip()[:30]
        
        # Puntos de captura optimizados para el flujo de pago
        candidates = [0.10, 0.25, 0.50, 0.75, 0.98] 
        
        for i, pos in enumerate(candidates):
            frame_idx = int(total_frames * pos)
            cap.set(cv2.CAP_PROP_POS_FRAMES, frame_idx)
            
            found_frame = False
            # Búsqueda rápida en un rango de 100 frames
            for attempt in range(100):
                ret, frame = cap.read()
                if not ret: break
                
                gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
                brightness = gray.mean()
                std_dev = gray.std()
                
                # Criterio equilibrado: Luz y Contraste básico
                if brightness > 45 and std_dev > 18:
                    frame_name = f"tmp_{video_stem.replace(' ', '_')}_{i+1}.jpg"
                    frame_path = os.path.join(output_folder, frame_name)
                    
                    success, img_encoded = cv2.imencode('.jpg', frame)
                    if success:
                        with open(frame_path, 'wb') as f:
                            f.write(img_encoded)
                        extracted_paths.append(frame_path)
                        print(f"[EvidenceService] Captura {i+1} OK.")
                        found_frame = True
                    break
            
            if not found_frame:
                # Fallback rápido si no hay nada ideal cerca
                cap.set(cv2.CAP_PROP_POS_FRAMES, frame_idx)
                ret, frame = cap.read()
                if ret:
                    frame_name = f"tmp_{video_stem.replace(' ', '_')}_{i+1}.jpg"
                    frame_path = os.path.join(output_folder, frame_name)
                    cv2.imwrite(frame_path, frame)
                    extracted_paths.append(frame_path)
        
        cap.release()
        return extracted_paths
    def generate_report(self, root_dir, output_name=None):
        """Escanea el directorio y genera el documento Word con soporte para archivos en raíz y subcarpetas."""
        root_dir = os.path.normpath(root_dir.strip().replace('"', '').replace("'", ""))
        
        if not os.path.exists(root_dir):
            raise Exception(f"La ruta especificada no existe: {root_dir}")
        
        # Generar nombre con timestamp
        timestamp = datetime.datetime.now().strftime("%d%m_%H%M")
        if output_name is None:
            output_name = f"Reporte_Evidencias_{timestamp}.docx"
            
        output_path = os.path.join(root_dir, output_name)

        # ESTRATEGIA: Copiar el template físicamente
        if os.path.exists(self.template_path):
            try:
                shutil.copy2(self.template_path, output_path)
                doc = Document(output_path)
                # Añadir info de cabecera después de la carátula
                doc.add_page_break()
                doc.add_heading(f"Reporte de Evidencias - {os.path.basename(root_dir)}", 1)
                doc.add_paragraph(f"Fecha: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}")
                doc.add_paragraph(f"Iniciativa: {os.path.basename(root_dir)}")
                doc.add_paragraph(f"Ruta: {root_dir}")
                doc.add_paragraph("-" * 50)
            except Exception as e:
                print(f"[EvidenceService] Error al copiar template: {str(e)}")
                doc = Document()
        else:
            doc = Document()
            doc.add_heading('Reporte de Evidencias de QA', 0)
            doc.add_paragraph(f"Fecha: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}")

        # 1. Procesar archivos sueltos en la raíz
        self._add_folder_to_doc(doc, root_dir, "Evidencias Generales (Raíz)", recursive=False)

        # 2. Procesar subcarpetas
        for item in sorted(os.listdir(root_dir)):
            item_path = os.path.join(root_dir, item)
            if os.path.isdir(item_path) and not item.startswith('_') and not item.startswith('.'):
                self._add_folder_to_doc(doc, item_path, item, recursive=True)

        try:
            doc.save(output_path)
        except PermissionError:
            raise Exception(f"No se pudo guardar el reporte porque el archivo está abierto. Por favor ciérralo.")
        
        # Limpieza sincronizada
        temp_dir = self._get_temp_dir()
        if os.path.exists(temp_dir):
            import time
            time.sleep(1)
            shutil.rmtree(temp_dir, ignore_errors=True)
            print(f"[EvidenceService] Carpeta temporal eliminada: {temp_dir}")

        return output_path

    def generate_html_report(self, root_dir, output_name=None):
        """Escanea el directorio y genera un reporte HTML con videos incrustados nativamente."""
        import base64
        root_dir = os.path.normpath(root_dir.strip().replace('"', '').replace("'", ""))
        
        if not os.path.exists(root_dir):
            raise Exception(f"La ruta especificada no existe: {root_dir}")
        
        timestamp = datetime.datetime.now().strftime("%d%m_%H%M")
        if output_name is None:
            output_name = f"Reporte_Evidencias_{timestamp}.html"
            
        output_path = os.path.join(root_dir, output_name)
        
        # Iniciar HTML con CSS moderno
        html_content = f"""
        <!DOCTYPE html>
        <html lang="es">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Reporte de Evidencias - {os.path.basename(root_dir)}</title>
            <style>
                :root {{
                    --bg-color: #0f172a;
                    --text-color: #f8fafc;
                    --card-bg: #1e293b;
                    --accent: #3b82f6;
                    --border: #334155;
                }}
                body {{
                    font-family: 'Segoe UI', system-ui, sans-serif;
                    background-color: var(--bg-color);
                    color: var(--text-color);
                    margin: 0;
                    padding: 2rem;
                    line-height: 1.6;
                }}
                .container {{
                    max-width: 1000px;
                    margin: 0 auto;
                }}
                .header {{
                    background: var(--card-bg);
                    padding: 2rem;
                    border-radius: 12px;
                    border: 1px solid var(--border);
                    margin-bottom: 2rem;
                    box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                }}
                h1 {{ margin-top: 0; color: var(--accent); font-size: 1.8rem; }}
                .section {{
                    background: var(--card-bg);
                    padding: 2rem;
                    border-radius: 12px;
                    border: 1px solid var(--border);
                    margin-bottom: 2rem;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                }}
                .section h2 {{
                    border-bottom: 2px solid var(--border);
                    padding-bottom: 0.5rem;
                    margin-top: 0;
                    color: #94a3b8;
                    font-size: 1.4rem;
                }}
                .media-item {{
                    margin-bottom: 2rem;
                    text-align: center;
                }}
                img, video {{
                    max-width: 100%;
                    max-height: 700px;
                    border-radius: 8px;
                    box-shadow: 0 4px 15px rgba(0,0,0,0.3);
                    background: #000;
                }}
                @media (max-width: 600px) {{
                    body {{ padding: 1rem; }}
                    .header, .section {{ padding: 1.2rem; }}
                    h1 {{ font-size: 1.5rem; }}
                    .section h2 {{ font-size: 1.2rem; }}
                }}
            </style>
        </head>
        <body>
            <div class="container">
                <div class="header">
                    <h1>Reporte de Evidencias de QA</h1>
                    <p><strong>Iniciativa:</strong> {os.path.basename(root_dir)}</p>
                    <p><strong>Fecha:</strong> {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}</p>
                    <p><strong>Ruta Original:</strong> {root_dir}</p>
                </div>
        """

        image_extensions = ('.png', '.jpg', '.jpeg', '.bmp', '.gif')
        video_extensions = ('.mp4', '.avi', '.mov', '.mkv', '.webm')

        def process_folder(folder_path, title):
            nonlocal html_content
            files = sorted(os.listdir(folder_path))
            evidence_files = [f for f in files if f.lower().endswith(image_extensions + video_extensions)]
            
            if not evidence_files:
                return

            html_content += f'<div class="section">\n<h2>{title}</h2>\n'
            
            for file in evidence_files:
                file_path = os.path.join(folder_path, file)
                
                try:
                    with open(file_path, "rb") as f:
                        file_data = f.read()
                    
                    b64_data = base64.b64encode(file_data).decode('utf-8')
                    
                    if file.lower().endswith(image_extensions):
                        ext = file.split('.')[-1].lower()
                        mime = f"image/{ext}" if ext != 'jpg' else "image/jpeg"
                        html_content += f"""
                        <div class="media-item">
                            <img src="data:{mime};base64,{b64_data}" alt="{file}" />
                        </div>
                        """
                    elif file.lower().endswith(video_extensions):
                        ext = file.split('.')[-1].lower()
                        mime = f"video/{ext}" if ext != 'mkv' else "video/x-matroska"
                        html_content += f"""
                        <div class="media-item">
                            <video controls preload="metadata">
                                <source src="data:{mime};base64,{b64_data}" type="{mime}">
                                Tu navegador no soporta el tag de video.
                            </video>
                        </div>
                        """
                except Exception as e:
                    print(f"[EvidenceService] ERROR al procesar para HTML {file}: {str(e)}")
                    html_content += f'<p style="color: #ef4444;">Error al cargar {file}</p>'
                    
            html_content += '</div>\n'

        # 1. Procesar raíz
        process_folder(root_dir, "Evidencias Generales (Raíz)")

        # 2. Procesar subcarpetas
        for item in sorted(os.listdir(root_dir)):
            item_path = os.path.join(root_dir, item)
            if os.path.isdir(item_path) and not item.startswith('_') and not item.startswith('.'):
                process_folder(item_path, item)

        html_content += """
            </div>
        </body>
        </html>
        """

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        print(f"[EvidenceService] Reporte HTML generado en: {output_path}")
        return output_path

    def _add_folder_to_doc(self, doc, folder_path, title, recursive=False):
        """Agrega el contenido de una carpeta al documento."""
        files = sorted(os.listdir(folder_path))
        image_extensions = ('.png', '.jpg', '.jpeg', '.bmp')
        video_extensions = ('.mp4', '.avi', '.mov', '.mkv')

        # Filtrar solo archivos de imagen/video para saber si añadir el título
        evidence_files = [f for f in files if f.lower().endswith(image_extensions + video_extensions)]
        
        if not evidence_files:
            return

        print(f"[EvidenceService] Añadiendo sección: {title}")
        doc.add_heading(title, level=1)
        
        files = sorted(os.listdir(folder_path))
        image_extensions = ('.png', '.jpg', '.jpeg', '.bmp')
        video_extensions = ('.mp4', '.avi', '.mov', '.mkv')

        found_content = False
        for file in files:
            file_path = os.path.join(folder_path, file)
            
            if file.lower().endswith(image_extensions):
                found_content = True
                try:
                    # Usar un ancho algo menor para asegurar que quepa en los márgenes
                    doc.add_picture(file_path, width=Inches(5.8))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"[EvidenceService] ERROR al insertar imagen {file}: {str(e)}")
                    doc.add_paragraph(f"(Error al insertar imagen)")
                doc.add_paragraph("") 

            elif file.lower().endswith(video_extensions):
                found_content = True
                # Las capturas ahora van a la carpeta temporal central
                frame_paths = self.extract_frames_from_video(file_path)
                
                if not frame_paths:
                    print(f"[EvidenceService] WARN: No se pudieron obtener capturas del video {file}")
                    doc.add_paragraph("(No se pudieron extraer capturas de este video)")
                
                for fp in frame_paths:
                    try:
                        print(f"[EvidenceService] Insertando captura de video en Word: {fp}")
                        # Usar prefijo para rutas largas si es necesario
                        docx_fp = fp
                        if len(docx_fp) > 240 and os.name == 'nt':
                            docx_fp = "\\\\?\\" + docx_fp
                            
                        doc.add_picture(docx_fp, width=Inches(5.8))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        doc.add_paragraph("")
                    except Exception as e:
                        print(f"[EvidenceService] ERROR al insertar captura de video {fp}: {str(e)}")

        if found_content:
            doc.add_page_break()

    def create_evidence_structure(self, initiative_name: str, test_cases: list):
        """Crea la estructura de carpetas de evidencia en Documents/Certificaciones."""
        import re
        
        # 1. Obtener ruta a Documentos/Certificaciones dinámicamente
        docs_path = os.path.expanduser("~\\Documents\\Certificaciones")
        
        # Limpiar nombre de la iniciativa para que sea válido en Windows
        safe_initiative = re.sub(r'[<>:"/\\|?*]', '_', initiative_name).strip()
        initiative_path = os.path.join(docs_path, safe_initiative)
        
        # Crear carpeta padre
        os.makedirs(initiative_path, exist_ok=True)
        
        # Crear subcarpetas
        created_folders = []
        for tc in test_cases:
            if not tc.strip(): continue
            safe_tc = re.sub(r'[<>:"/\\|?*]', '_', tc).strip()
            tc_path = os.path.join(initiative_path, safe_tc)
            os.makedirs(tc_path, exist_ok=True)
            created_folders.append(safe_tc)
            
        # Abrir la carpeta en el explorador de archivos de Windows
        if os.name == 'nt':
            os.startfile(initiative_path)
            
        return {
            "path": initiative_path,
            "folders_created": len(created_folders)
        }
